"""
=============================================================================
  LIFE SIMULATOR — GUI Edition  (pygame)  -  Estadísticas mejoradas
=============================================================================
  Ventana gráfica con estética RPG cinemático oscuro.
  Requiere:  pip install pygame matplotlib networkx
=============================================================================
"""

import math, os, random, sqlite3, sys, textwrap, time
from dataclasses import dataclass, field
from datetime import datetime

import pygame
import pygame.gfxdraw

try:
    import docx as _docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    HAS_MPL = True
except ImportError:
    HAS_MPL = False

try:
    import networkx as nx
    HAS_NX = True
except ImportError:
    HAS_NX = False

# ══════════════════════════════════════════════════════════════════════════════
#  CONSTANTES
# ══════════════════════════════════════════════════════════════════════════════

W, H        = 1100, 720
FPS         = 60
DB_PATH     = "life_sim_achievements.db"
GRAPH_OUT   = "life_graph.png"

# Paleta
C = {
    "bg":          (10,  10,  18),
    "panel":       (18,  18,  32),
    "panel2":      (24,  24,  44),
    "border":      (55,  55,  90),
    "border_hi":   (120, 100, 200),
    "gold":        (212, 175, 55),
    "gold_dim":    (120, 100, 35),
    "text":        (220, 215, 240),
    "text_dim":    (120, 115, 140),
    "accent":      (140, 90,  255),
    "accent2":     (80,  200, 255),
    "green":       (80,  220, 120),
    "red":         (220, 70,  70),
    "red_dim":     (120, 40,  40),
    "orange":      (255, 160, 50),
    "white":       (255, 255, 255),
    "stat_bars": {
        "salud":        (80,  220, 120),
        "inteligencia": (80,  180, 255),
        "carisma":      (255, 150, 200),
        "valentia":     (255, 120, 50),
        "dinero":       (212, 175, 55),
        "suerte":       (180, 130, 255),
    },
}

DIFFICULTY_SETTINGS = {
    "Europa":        {"label": "Muy Fácil",   "risk_mult": 0.3,  "emoji": "🌍"},
    "Norteamérica":  {"label": "Fácil",       "risk_mult": 0.5,  "emoji": "🌎"},
    "Oceanía":       {"label": "Medio",       "risk_mult": 0.7,  "emoji": "🌏"},
    "Asia":          {"label": "Difícil",     "risk_mult": 1.0,  "emoji": "🗺"},
    "Latinoamérica": {"label": "Muy Difícil", "risk_mult": 1.3,  "emoji": "🌎"},
    "África":        {"label": "Hardcore",    "risk_mult": 1.8,  "emoji": "🌍"},
}

ALL_TITLES = [
    "El Ícono", "El Sabio", "El Justo", "El Patriarca", "El Redimido",
    "El Fantasma", "El Rentista", "El Magnate", "El Estadista",
    "El Tranquilo", "El Equilibrado", "El Filántropo", "El Mentor",
    "El Pensador", "El Maldito", "El Fugitivo", "El Caído",
    "El Héroe Oscuro", "El Explorador", "El Trabajador", "El Mártir",
    "El Temerario",
]


# ══════════════════════════════════════════════════════════════════════════════
#  ARCHIVOS DE PREGUNTAS POR DIFICULTAD
# ══════════════════════════════════════════════════════════════════════════════

import re as _re

_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__)) if '__file__' in dir() else os.getcwd()

QUESTION_FILES = {
    "Europa":        os.path.join(_SCRIPT_DIR, "Preguntas Europa.docx"),
    "Norteamérica":  os.path.join(_SCRIPT_DIR, "Preguntas Norte America.docx"),
    "Oceanía":       os.path.join(_SCRIPT_DIR, "Preguntas Oceania.docx"),
    "Asia":          os.path.join(_SCRIPT_DIR, "preguntas_Asia.sql"),
    "Latinoamérica": os.path.join(_SCRIPT_DIR, "Preguntas Latinoamerica.docx"),
    "África":        os.path.join(_SCRIPT_DIR, "Preguntas Africa.docx"),
}

_STAT_MAP = {
    "Inteligencia": "inteligencia", "Salud": "salud",
    "Carisma": "carisma", "Valentía": "valentia",
    "Valentia": "valentia", "Dinero": "dinero", "Suerte": "suerte",
}


def _parse_stats(text):
    r = {}
    for val, stat in _re.findall(
            r'([+-]\d+)\s+(Inteligencia|Salud|Carisma|Valentía|Valentia|Dinero|Suerte)', text):
        r[_STAT_MAP[stat]] = int(val)
    for stat, val in _re.findall(
            r'(Inteligencia|Salud|Carisma|Valentía|Valentia|Dinero|Suerte)\s+([+-]?\d+)', text):
        k = _STAT_MAP[stat]
        if k not in r:
            r[k] = int(val)
    return r


def _parse_europa_questions(file_path):
    doc = _docx.Document(file_path)
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    SKIP = {'INFANCIA', 'ADOLESCENCIA', 'JUVENTUD', 'MADUREZ', 'VEJEZ',
            'ADULTEZ', 'TRABAJO', 'UNIVERSIDAD', 'VIDA ADULTA', 'EDAD ADULTA'}
    Q_PAT   = _re.compile(r'^(\d+[-]?\d*)[\.:]?\s+(.+)')
    OPT_PAT = _re.compile(r'^\s*[•\-]?\s*(.+?)\s*[→:]\s*(.+)')
    questions = []
    i = 0
    while i < len(paras):
        p = paras[i]
        upper = p.upper().strip()
        if upper in SKIP or upper.startswith('FASE') or p.startswith('\U0001f6e0') or p.startswith('\U0001f3db'):
            i += 1; continue
        m = Q_PAT.match(p)
        if m:
            qtext = m.group(2).strip()
            opts = []
            j = i + 1
            while j < len(paras) and len(opts) < 2:
                np = paras[j]
                nu = np.upper().strip()
                if nu in SKIP or nu.startswith('FASE') or np.startswith('\U0001f6e0') or np.startswith('\U0001f3db'):
                    j += 1; continue
                if Q_PAT.match(np):
                    break
                om = OPT_PAT.match(np)
                if om:
                    opts.append((om.group(1).strip()[:80], _parse_stats(om.group(2))))
                j += 1
            if len(opts) == 2:
                questions.append((qtext, opts[0][0], opts[0][1], opts[1][0], opts[1][1]))
                i = j
            else:
                i += 1
        else:
            i += 1
    return questions


def _parse_continental_docx(file_path):
    doc = _docx.Document(file_path)
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    Q_PAT = _re.compile(r'^(\d+)\.\s*(.+)')
    OPT_A = _re.compile(r'^A:\s*(.+?)\s*->\s*[Vv]e\s+a\s+la\s+(\d+)')
    OPT_B = _re.compile(r'^B:\s*(.+?)\s*->\s*[Vv]e\s+a\s+la\s+(\d+)')
    REDIR = _re.compile(r'->\s*[Vv]e\s+a\s+la\s+(\d+)')
    raw = {}
    num_to_nid = {}
    i = 0
    while i < len(paras):
        p = paras[i]
        m = Q_PAT.match(p)
        if not m:
            i += 1; continue
        qnum  = int(m.group(1))
        qtext = m.group(2).strip()
        qtext_full = qtext
        age   = 5 + int(qnum / 100 * 70)
        final_m = _re.search(r'\(FINAL[:\s]+([^)]+)\)', qtext)
        go_m    = _re.search(r'\(GAME OVER[:\s]*([^)]*)\)', qtext)
        bucle_m = 'BUCLE' in qtext.upper()
        is_redir = 'SALTO NARRATIVO' in qtext or 'NODO ELIMINADO' in qtext
        redir_m  = REDIR.search(qtext)

        if final_m or go_m or bucle_m:
            raw_title = (final_m.group(1) if final_m else (go_m.group(1) if go_m else "El Bucle"))
            short    = raw_title.split(':')[-1].strip().split('.')[0].strip()[:30]
            is_death = bool(go_m)
            if bucle_m:
                nid    = f"bucle_{qnum}"
                opts_b = []
                j = i + 1
                while j < min(i + 4, len(paras)):
                    ma = OPT_A.match(paras[j])
                    mb = OPT_B.match(paras[j])
                    if ma:
                        opts_b.append({'text': ma.group(1)[:150], 'next_num': None, 'next': f"old_{qnum}_fin"})
                        j += 1
                    elif mb:
                        opts_b.append({'text': mb.group(1)[:150], 'next_num': 1, 'next': 'q_1'})
                        j += 1
                    else:
                        break
                fin_nid = f"old_{qnum}_fin"
                raw[f"fin_{qnum}"] = {'nid': fin_nid, 'terminal': True, 'age': age + 5,
                    'title': "El Superviviente", 'death_msg': "Elegiste la esperanza. Tu historia sigue adelante.",
                    'is_death': False}
                raw[qnum] = {'nid': nid, 'terminal': False, 'age': age, 'situation': qtext_full, 'opts': opts_b}
                num_to_nid[qnum] = nid
                i = j
            else:
                nid = f"death_{qnum}" if is_death else f"old_{qnum}"
                _cm = _re.sub(r'^\s*\((?:FINAL|GAME OVER)[^)]*\)\s*\.?\s*','',qtext_full).strip()
                raw[qnum] = {'nid': nid, 'terminal': True, 'age': age,
                    'title': short, 'death_msg': _cm or qtext_full, 'is_death': is_death}
                num_to_nid[qnum] = nid
                i += 1
        elif is_redir and redir_m:
            next_num = int(redir_m.group(1))
            raw[qnum] = {'nid': f"q_{qnum}", 'terminal': False, 'age': age, 'situation': "...",
                'opts': [{'text': "Continuar", 'next_num': next_num}]}
            num_to_nid[qnum] = f"q_{qnum}"
            i += 1
        else:
            # Recoger párrafos de continuación antes de A:/B:
            j = i + 1
            extra = []
            while j < len(paras):
                np = paras[j]
                if Q_PAT.match(np) or OPT_A.match(np) or OPT_B.match(np):
                    break
                extra.append(np)
                j += 1
            if extra:
                qtext_full = qtext + ' ' + ' '.join(extra)

            # Detectar redirect inline (-> Ve a la N) ANTES de buscar A/B
            _inline_redir = REDIR.search(qtext_full)
            if _inline_redir:
                # Limpiar el "-> Ve a la N" del texto de situación
                qtext_full = qtext_full[:_inline_redir.start()].strip().rstrip('.')

            opts = []
            while j < len(paras):
                ma = OPT_A.match(paras[j])
                mb = OPT_B.match(paras[j])
                if ma:
                    opts.append({'text': ma.group(1)[:150], 'next_num': int(ma.group(2))})
                    j += 1
                elif mb:
                    opts.append({'text': mb.group(1)[:150], 'next_num': int(mb.group(2))})
                    j += 1
                else:
                    break

            # Si no hay opciones A/B pero hay redirect inline, usar como único camino
            if not opts and _inline_redir:
                opts = [{'text': "Continuar", 'next_num': int(_inline_redir.group(1))}]

            raw[qnum] = {'nid': f"q_{qnum}", 'terminal': False, 'age': age,
                'situation': qtext_full, 'opts': opts}
            num_to_nid[qnum] = f"q_{qnum}"
            i = j

    def resolve(nn):
        if nn is None: return None
        return num_to_nid.get(nn, f"q_{nn}")

    nodes = {}
    for key, data in raw.items():
        nid = data['nid']
        if data.get('terminal'):
            nodes[nid] = data
        else:
            resolved = [{'text': o['text'], 'next': o.get('next') or resolve(o.get('next_num'))}
                        for o in data.get('opts', [])]
            nodes[nid] = {**data, 'opts': resolved}

    fallback_nid = next(
        (nid for nid in nodes if nodes[nid].get('terminal') and not nodes[nid].get('is_death')),
        next((nid for nid in nodes if nodes[nid].get('terminal')), None))
    for nid, nd in list(nodes.items()):
        if nd.get('terminal'): continue
        for o in nd.get('opts', []):
            if o['next'] not in nodes and fallback_nid:
                o['next'] = fallback_nid
    return nodes


def _parse_asia_sql(sql_path):
    with open(sql_path, encoding='utf-8') as f:
        content = f.read()
    TERM_RE = _re.compile(
        r"\((\d+),\s*'((?:[^'\\]|\\.)*)',\s*NULL,\s*NULL,\s*NULL,\s*NULL\)", _re.DOTALL)
    ROW_RE  = _re.compile(
        r"\((\d+),\s*'((?:[^'\\]|\\.)*)',\s*'((?:[^'\\]|\\.)*)',"
        r"\s*'((?:[^'\\]|\\.)*)',\s*(\d+),\s*(\d+)\)", _re.DOTALL)
    num_to_nid = {}
    nodes = {}
    for m in TERM_RE.finditer(content):
        qnum = int(m.group(1))
        q    = m.group(2)
        final_m = _re.search(r'\[FINAL[:\s]+([^\]]+)\]', q)
        go_m    = _re.search(r'\[GAME OVER[:\s]*([^\]]*)\]', q)
        is_death = bool(go_m)
        raw  = (final_m.group(1) if final_m else (go_m.group(1) if go_m else ""))
        short = raw.split(':')[-1].strip()[:30]
        nid  = f"death_{qnum}" if is_death else f"old_{qnum}"
        num_to_nid[qnum] = nid
        _cq = _re.sub(r'^\s*\[(?:FINAL|GAME OVER)[^\]]*\]\s*\.?\s*','',q).strip()
        nodes[nid] = {'terminal': True, 'age': 70, 'title': short,
                      'death_msg': _cq or q, 'is_death': is_death}
    for m in ROW_RE.finditer(content):
        qnum = int(m.group(1))
        q  = m.group(2)
        a  = m.group(3)
        b  = m.group(4)
        na = int(m.group(5))
        nb = int(m.group(6))
        age_m = _re.search(r'\[Edad:\s*(\d+)', q)
        age   = int(age_m.group(1)) if age_m else 5 + int(qnum / 200 * 70)
        nodes[f"q_{qnum}"] = {
            'terminal': False, 'age': age, 'situation': q,
            'opts': [{'text': a[:150], 'next': num_to_nid.get(na, f"q_{na}")},
                     {'text': b[:150], 'next': num_to_nid.get(nb, f"q_{nb}")}]}
    return nodes



def _heuristic_stats(opt_text, situation=""):
    txt = (opt_text + " " + situation).lower()
    s = {}
    def add(k, v): s[k] = s.get(k, 0) + v
    if any(w in txt for w in ['estudi','aprend','leer','libro','universid','investig','debate']): add('inteligencia', 1)
    if any(w in txt for w in ['descans','dorm','médic','hospital','deport','ejercic','curar']): add('salud', 1)
    if any(w in txt for w in ['droga','alcohol','herida','golpe','veneno','enferm','hambre']): add('salud', -1)
    if any(w in txt for w in ['trabaj','gana','cobr','sueldo','empleo','negoci','vend','ahorra']): add('dinero', 1)
    if any(w in txt for w in ['pagar','deuda','multa','pierd','quiebra']): add('dinero', -1)
    if any(w in txt for w in ['amig','famili','ayud','comunid','colabor','social','pareja','amor']): add('carisma', 1)
    if any(w in txt for w in ['traicion','mentir','engañ','huy','escond','aislar','rechazo']): add('carisma', -1)
    if any(w in txt for w in ['luch','enfrent','resiste','defien','arriesg','atac','rebel']): add('valentia', 1)
    if any(w in txt for w in ['rend','agacho','obedezco','suplic','huyo','escondo','callo']): add('valentia', -1)
    if any(w in txt for w in ['suert','azar','oport','encontr','hall']): add('suerte', 1)
    if any(w in txt for w in ['rob','estaf','corrup','traf','ilegal','crimen','pandilla','asesin']):
        add('dinero', 1); add('inteligencia', -1); add('carisma', -1)
    return {k: v for k, v in s.items() if v != 0}

def _to_decision_nodes(raw_nodes, start_key="q_1"):
    nodes = {}
    for nid, data in raw_nodes.items():
        if data.get('terminal'):
            node = DecisionNode(
                node_id=nid, age=data.get('age', 70),
                situation="Ha llegado tu momento...",
                options=[], is_terminal=True,
                death_msg=data.get('death_msg', ''),
                title=data.get('title', ''))
        else:
            _sit = data.get('situation','')
            options = [
                {'text': o['text'],
                 'stat_delta': _heuristic_stats(o['text'], _sit),
                 'next_node': o['next'], 'death_risk': 0.0}
                for o in data.get('opts', []) if o.get('next')]
            node = DecisionNode(
                node_id=nid, age=data.get('age', 20),
                situation=data.get('situation', ''),
                options=options)
        nodes[nid] = node
    if start_key in nodes:
        nodes["start"] = nodes[start_key]
    elif nodes:
        nodes["start"] = next(iter(nodes.values()))
    return nodes



def _collapse_redirects(nodes):
    def follow(nid, seen=None):
        if seen is None: seen = set()
        if nid in seen or nid not in nodes: return nid
        seen.add(nid); nd = nodes[nid]
        if (not nd.is_terminal and len(nd.options)==1
                and nd.options[0]['text']=='Continuar' and nd.situation=='...'):
            return follow(nd.options[0]['next_node'], seen)
        return nid
    for nd in nodes.values():
        if nd.is_terminal: continue
        for opt in nd.options:
            opt['next_node'] = follow(opt['next_node'])
    to_del = [nid for nid,nd in nodes.items()
              if nid!='start' and not nd.is_terminal
              and len(nd.options)==1 and nd.options[0]['text']=='Continuar'
              and nd.situation=='...']
    for nid in to_del: del nodes[nid]
    return nodes

def _build_europa_tree(file_path):
    if not os.path.exists(file_path):
        return None
    questions = _parse_europa_questions(file_path)
    if not questions:
        return None
    nodes = {}
    total = len(questions)
    for idx, (qtext, tA, sA, tB, sB) in enumerate(questions):
        age     = 5 + int(idx / total * 75)
        nid     = "start" if idx == 0 else f"q_{idx}"
        next_nid = f"q_{idx + 1}" if idx < total - 1 else "old_europa_end"
        nodes[nid] = DecisionNode(
            node_id=nid, age=age, situation=qtext,
            options=[
                {'text': tA, 'stat_delta': sA, 'next_node': next_nid, 'death_risk': 0.0},
                {'text': tB, 'stat_delta': sB, 'next_node': next_nid, 'death_risk': 0.0},
            ])
    nodes["old_europa_end"] = DecisionNode(
        node_id="old_europa_end", age=80,
        situation="Llegas a los 80 años. Es el momento de partir.",
        options=[], is_terminal=True,
        death_msg="Has tomado tus propias decisiones y vivido con sus consecuencias. Una vida tranquila y completa en Europa.",
        title="El Europeo Sereno")
    return nodes


def _build_continental_tree(file_path):
    if not os.path.exists(file_path):
        return None
    raw = _parse_continental_docx(file_path)
    if not raw:
        return None
    nodes = _to_decision_nodes(raw, start_key="q_1")
    return _collapse_redirects(nodes)


def _build_asia_tree(file_path):
    if not os.path.exists(file_path):
        return None
    raw = _parse_asia_sql(file_path)
    if not raw:
        return None
    nodes = _to_decision_nodes(raw, start_key="q_1")
    return _collapse_redirects(nodes)


def build_difficulty_tree(difficulty):
    """Construye el árbol de decisiones según la dificultad. Fallback al árbol original si falla."""
    if not HAS_DOCX and difficulty != "Asia":
        return build_decision_tree()
    try:
        fpath = QUESTION_FILES.get(difficulty, "")
        if difficulty == "Europa":
            tree = _build_europa_tree(fpath)
        elif difficulty == "Asia":
            tree = _build_asia_tree(fpath)
        else:
            tree = _build_continental_tree(fpath)
        if tree and "start" in tree:
            return tree
    except Exception as e:
        print(f"[WARN] Error cargando preguntas para {difficulty}: {e}")
    return build_decision_tree()


# ══════════════════════════════════════════════════════════════════════════════
#  ÁRBOL DE DECISIONES  (idéntico al original)
# ══════════════════════════════════════════════════════════════════════════════

@dataclass
class DecisionNode:
    node_id:     str
    age:         int
    situation:   str
    options:     list = field(default_factory=list)
    is_terminal: bool = False
    death_msg:   str  = ""
    title:       str  = ""

def build_decision_tree():
    nodes = {}
    def add(nid, age, sit, opts, terminal=False, dmsg="", title=""):
        nodes[nid] = DecisionNode(nid, age, sit, opts, terminal, dmsg, title)

    add("start", 5,
        "Tienes 5 años. Un perro grande se acerca ladrando. ¿Qué haces?",
        [{"text":"Te quedas quieto y lo miras a los ojos",
          "stat_delta":{"valentia":+2,"inteligencia":+1},"next_node":"child_brave","death_risk":0.0},
         {"text":"Sales corriendo hacia tu madre",
          "stat_delta":{"carisma":+1,"salud":-1},"next_node":"child_scared","death_risk":0.05}])

    add("child_brave", 8,
        "A los 8 años descubres una cueva en el bosque. ¿Entras?",
        [{"text":"Entras con una antorcha improvisada",
          "stat_delta":{"valentia":+2,"suerte":+1},"next_node":"teen_explorer","death_risk":0.05},
         {"text":"Vuelves a casa a buscar un adulto",
          "stat_delta":{"inteligencia":+2,"carisma":+1},"next_node":"teen_studious","death_risk":0.0}])

    add("child_scared", 8,
        "Con 8 años te invitan al equipo de fútbol del barrio. ¿Aceptas?",
        [{"text":"Aceptas aunque te da vergüenza",
          "stat_delta":{"carisma":+2,"salud":+1},"next_node":"teen_social","death_risk":0.0},
         {"text":"Rechazas y prefieres leer en casa",
          "stat_delta":{"inteligencia":+3},"next_node":"teen_studious","death_risk":0.0}])

    add("teen_explorer", 14,
        "14 años. Tus amigos proponen escalar el tejado del colegio de noche.",
        [{"text":"Subes el primero, sin miedo",
          "stat_delta":{"valentia":+2,"salud":-1},"next_node":"young_rebel","death_risk":0.08},
         {"text":"Propones explorar las alcantarillas en su lugar",
          "stat_delta":{"inteligencia":+1,"carisma":+1},"next_node":"young_clever","death_risk":0.02}])

    add("teen_studious", 14,
        "Con 14 años encuentras un libro de filosofía en el desván. ¿Qué haces?",
        [{"text":"Lo lees entero en una semana",
          "stat_delta":{"inteligencia":+3,"suerte":+1},"next_node":"young_scholar","death_risk":0.0},
         {"text":"Lo dejas a medias y sales a jugar",
          "stat_delta":{"carisma":+2},"next_node":"young_social_teen","death_risk":0.0}])

    add("teen_social", 14,
        "A los 14 tu equipo llega a la final regional. El día del partido tienes fiebre.",
        [{"text":"Juegas igual; el equipo te necesita",
          "stat_delta":{"carisma":+2,"salud":-2},"next_node":"young_athlete","death_risk":0.07},
         {"text":"Te quedas en casa y el equipo pierde",
          "stat_delta":{"inteligencia":+1,"salud":+1},"next_node":"young_responsible","death_risk":0.0}])

    add("young_rebel", 19,
        "Tienes 19 años. Te ofrecen tocar en una banda underground o estudiar Derecho.",
        [{"text":"La música: vives el sueño",
          "stat_delta":{"valentia":+2,"carisma":+2,"dinero":-2},"next_node":"adult_artist","death_risk":0.05},
         {"text":"Derecho: seguridad a largo plazo",
          "stat_delta":{"inteligencia":+2,"dinero":+2},"next_node":"adult_lawyer","death_risk":0.0}])

    add("young_clever", 19,
        "A los 19 descubres que tienes talento para hackear sistemas.",
        [{"text":"Ethical hacking: reportas fallos a empresas",
          "stat_delta":{"inteligencia":+3,"dinero":+1},"next_node":"adult_hacker_white","death_risk":0.0},
         {"text":"Robas pequeñas sumas de cuentas bancarias",
          "stat_delta":{"dinero":+4,"inteligencia":+1},"next_node":"adult_criminal","death_risk":0.15}])

    add("young_scholar", 19,
        "Con 19 te admiten en dos universidades: una local gratuita y una de élite con deudas.",
        [{"text":"Universidad de élite — prestas el dinero",
          "stat_delta":{"inteligencia":+3,"dinero":-3},"next_node":"adult_elite_grad","death_risk":0.0},
         {"text":"Universidad local — sin deudas",
          "stat_delta":{"inteligencia":+2,"dinero":+1},"next_node":"adult_humble_grad","death_risk":0.0}])

    add("young_social_teen", 19,
        "Con 19 años tus amigos te invitan a viajar por el mundo de forma improvisada.",
        [{"text":"Te unes sin pensarlo",
          "stat_delta":{"carisma":+3,"suerte":+2,"dinero":-2},"next_node":"adult_wanderer","death_risk":0.05},
         {"text":"Rechazas y empiezas a trabajar localmente",
          "stat_delta":{"dinero":+2,"inteligencia":+1},"next_node":"adult_worker","death_risk":0.0}])

    add("young_athlete", 19,
        "Con 19 un agente deportivo te ofrece un contrato semiprofesional.",
        [{"text":"Firmas y dedicas tu vida al deporte",
          "stat_delta":{"salud":+3,"carisma":+2,"dinero":+1},"next_node":"adult_athlete","death_risk":0.05},
         {"text":"Rechazas; el cuerpo no lo aguantaría",
          "stat_delta":{"inteligencia":+2,"salud":+1},"next_node":"adult_humble_grad","death_risk":0.0}])

    add("young_responsible", 19,
        "A los 19 eres el más sensato del grupo. Te proponen ser representante estudiantil.",
        [{"text":"Aceptas el cargo con entusiasmo",
          "stat_delta":{"carisma":+3,"inteligencia":+1},"next_node":"adult_politician","death_risk":0.0},
         {"text":"Rechazas; prefieres las sombras",
          "stat_delta":{"inteligencia":+2,"suerte":+1},"next_node":"adult_scholar_deep","death_risk":0.0}])

    add("adult_artist", 32,
        "Tienes 32 años. Tu banda está en la cima, pero aparece el alcohol y las drogas.",
        [{"text":"Te mantienes sobrio; el arte primero",
          "stat_delta":{"valentia":+1,"salud":+2},"next_node":"midlife_legend","death_risk":0.03},
         {"text":"Te dejas llevar; vives al máximo",
          "stat_delta":{"carisma":+1,"salud":-4},"next_node":"death_overdose","death_risk":0.55}])

    add("adult_lawyer", 32,
        "Con 32, un cliente corrupto te ofrece el doble si le ayudas a evadir impuestos.",
        [{"text":"Lo rechazas y lo denuncias",
          "stat_delta":{"inteligencia":+2,"carisma":+1},"next_node":"midlife_honorable","death_risk":0.0},
         {"text":"Aceptas; nadie lo sabrá",
          "stat_delta":{"dinero":+4,"inteligencia":-1},"next_node":"midlife_corrupt","death_risk":0.2}])

    add("adult_hacker_white", 32,
        "A los 32, una agencia de inteligencia quiere reclutarte para operaciones secretas.",
        [{"text":"Aceptas la misión de alto riesgo",
          "stat_delta":{"valentia":+3,"inteligencia":+2},"next_node":"midlife_spy","death_risk":0.18},
         {"text":"Rechazas y fundes tu empresa de ciberseguridad",
          "stat_delta":{"dinero":+3,"inteligencia":+2},"next_node":"midlife_entrepreneur","death_risk":0.0}])

    add("adult_criminal", 32,
        "Con 32 la policía está cada vez más cerca. ¿Qué haces?",
        [{"text":"Entregas a tus cómplices y negocias inmunidad",
          "stat_delta":{"inteligencia":+1,"carisma":-2},"next_node":"midlife_reformed","death_risk":0.1},
         {"text":"Huyes a otro país con el dinero",
          "stat_delta":{"suerte":+2,"dinero":+2},"next_node":"death_fugitive","death_risk":0.45}])

    add("adult_elite_grad", 32,
        "Tienes 32. Tu carrera despega, pero te ofrecen un puesto en el extranjero lejos de tu familia.",
        [{"text":"Aceptas; la carrera primero",
          "stat_delta":{"dinero":+3,"carisma":-1},"next_node":"midlife_expat","death_risk":0.0},
         {"text":"Rechazas y construyes tu vida en casa",
          "stat_delta":{"carisma":+2,"suerte":+1},"next_node":"midlife_family_man","death_risk":0.0}])

    add("adult_humble_grad", 32,
        "Con 32, tu empresa está al borde de la quiebra. ¿Arriesgas los ahorros o buscas empleo?",
        [{"text":"Arriesgas todo en un último proyecto",
          "stat_delta":{"valentia":+3,"dinero":+3},"next_node":"midlife_entrepreneur","death_risk":0.1},
         {"text":"Buscas empleo seguro en una multinacional",
          "stat_delta":{"dinero":+1,"inteligencia":+1},"next_node":"midlife_corporate","death_risk":0.0}])

    add("adult_wanderer", 32,
        "Con 32 llevas años viajando. Conoces a alguien que te pide que te quedes.",
        [{"text":"Te quedas por amor",
          "stat_delta":{"carisma":+2,"suerte":+2},"next_node":"midlife_family_man","death_risk":0.0},
         {"text":"Sigues viajando; eres libre",
          "stat_delta":{"valentia":+2,"suerte":+1},"next_node":"midlife_nomad","death_risk":0.05}])

    add("adult_worker", 32,
        "A los 32 te ofrecen un ascenso que requiere 80 horas semanales.",
        [{"text":"Aceptas; el sacrificio vale",
          "stat_delta":{"dinero":+4,"salud":-2},"next_node":"midlife_burnout","death_risk":0.1},
         {"text":"Rechazas y mantienes el equilibrio",
          "stat_delta":{"salud":+2,"carisma":+1},"next_node":"midlife_balanced","death_risk":0.0}])

    add("adult_athlete", 32,
        "Con 32 sufres una lesión grave. ¿Te operas (arriesgado) o te retiras (seguro)?",
        [{"text":"Te operas y vuelves a competir",
          "stat_delta":{"valentia":+3,"salud":+1},"next_node":"midlife_legend","death_risk":0.12},
         {"text":"Te retiras con dignidad",
          "stat_delta":{"salud":+2,"carisma":+2},"next_node":"midlife_coach","death_risk":0.0}])

    add("adult_politician", 32,
        "Tienes 32. El partido te pide comprometer tu ética para avanzar.",
        [{"text":"Te niegas y renuncias al partido",
          "stat_delta":{"carisma":+3,"inteligencia":+1},"next_node":"midlife_activist","death_risk":0.0},
         {"text":"Cedes; el poder merece el precio",
          "stat_delta":{"dinero":+3,"carisma":+1},"next_node":"midlife_corrupt","death_risk":0.1}])

    add("adult_scholar_deep", 32,
        "A los 32 escribes una tesis revolucionaria que desafía a la academia.",
        [{"text":"La publicas y afrontas la controversia",
          "stat_delta":{"inteligencia":+4,"carisma":+1},"next_node":"midlife_thinker","death_risk":0.0},
         {"text":"La guardas; el mundo no está listo",
          "stat_delta":{"inteligencia":+2,"suerte":+2},"next_node":"midlife_hermit","death_risk":0.0}])

    midlife = [
        ("midlife_legend", 45,
         "A los 45 eres una leyenda. Una biopic sobre tu vida: ¿la apruebas?",
         [{"text":"Apruebas el guion y te expones",
           "stat_delta":{"carisma":+3},"next_node":"old_icon","death_risk":0.0},
          {"text":"La rechazas; tu historia es tuya",
           "stat_delta":{"inteligencia":+2,"suerte":+1},"next_node":"old_sage","death_risk":0.0}]),
        ("midlife_honorable", 45,
         "Con 45 te proponen como juez de la Corte Suprema.",
         [{"text":"Aceptas el cargo vitalicio",
           "stat_delta":{"inteligencia":+3,"carisma":+2},"next_node":"old_judge","death_risk":0.0},
          {"text":"Rechazas para estar con los tuyos",
           "stat_delta":{"carisma":+3,"suerte":+2},"next_node":"old_family","death_risk":0.0}]),
        ("midlife_corrupt", 45,
         "Tienes 45. Tu red de corrupción se desmorona. Un periodista lo sabe todo.",
         [{"text":"Confiesas y pides clemencia",
           "stat_delta":{"carisma":-1,"inteligencia":+1},"next_node":"old_redeemed","death_risk":0.05},
          {"text":"Intentas silenciar al periodista",
           "stat_delta":{"dinero":+2,"valentia":-2},"next_node":"death_scandal","death_risk":0.5}]),
        ("midlife_spy", 45,
         "Con 45, tras décadas en las sombras, una misión final podría cambiar la historia.",
         [{"text":"Aceptas la misión suicida",
           "stat_delta":{"valentia":+4},"next_node":"death_hero","death_risk":0.6},
          {"text":"Desapareces con una identidad nueva",
           "stat_delta":{"inteligencia":+2,"suerte":+3},"next_node":"old_ghost","death_risk":0.0}]),
        ("midlife_entrepreneur", 45,
         "Tu empresa vale millones. ¿La vendes y te retiras o sigues expandiendo?",
         [{"text":"La vendes y te retiras joven",
           "stat_delta":{"dinero":+5,"salud":+2},"next_node":"old_retired_rich","death_risk":0.0},
          {"text":"Sigues expandiendo globalmente",
           "stat_delta":{"dinero":+3,"salud":-1},"next_node":"old_tycoon","death_risk":0.05}]),
        ("midlife_reformed", 45,
         "Tras redimirte, un documental quiere contar tu historia criminal.",
         [{"text":"Colaboras con total honestidad",
           "stat_delta":{"carisma":+3,"inteligencia":+1},"next_node":"old_redeemed","death_risk":0.0},
          {"text":"Rechazas; el pasado debe morir",
           "stat_delta":{"suerte":+2},"next_node":"old_quiet","death_risk":0.0}]),
        ("midlife_expat", 45,
         "Con 45, llevas años en el extranjero. Tu país te necesita.",
         [{"text":"Vuelves a casa a servir a tu nación",
           "stat_delta":{"carisma":+3,"inteligencia":+2},"next_node":"old_statesman","death_risk":0.05},
          {"text":"Te quedas; ya construiste tu vida aquí",
           "stat_delta":{"dinero":+2,"suerte":+1},"next_node":"old_quiet","death_risk":0.0}]),
        ("midlife_family_man", 45,
         "Con 45 tus hijos crecen y necesitas reinventarte.",
         [{"text":"Emprendes un negocio familiar",
           "stat_delta":{"carisma":+2,"dinero":+2},"next_node":"old_family","death_risk":0.03},
          {"text":"Vuelves a estudiar una nueva carrera",
           "stat_delta":{"inteligencia":+3},"next_node":"old_sage","death_risk":0.0}]),
        ("midlife_nomad", 45,
         "A los 45 una enfermedad tropical te detiene en mitad de un viaje remoto.",
         [{"text":"Te automédicas y sigues",
           "stat_delta":{"valentia":+1,"salud":-3},"next_node":"death_illness","death_risk":0.5},
          {"text":"Buscas ayuda médica urgente",
           "stat_delta":{"salud":+1,"inteligencia":+1},"next_node":"old_quiet","death_risk":0.08}]),
        ("midlife_burnout", 45,
         "Tienes 45 y colapso total: corazón, familia, todo se derrumba.",
         [{"text":"Pides ayuda y tomas un año sabático",
           "stat_delta":{"salud":+3,"carisma":+1},"next_node":"old_balanced","death_risk":0.05},
          {"text":"Ignoras las señales y sigues trabajando",
           "stat_delta":{"dinero":+1,"salud":-4},"next_node":"death_heart","death_risk":0.65}]),
        ("midlife_balanced", 45,
         "Con 45 tienes estabilidad. ¿La usas para algo mayor?",
         [{"text":"Creas una fundación benéfica",
           "stat_delta":{"carisma":+3,"suerte":+2},"next_node":"old_philanthropist","death_risk":0.0},
          {"text":"Disfrutas de la calma ganada",
           "stat_delta":{"salud":+2,"suerte":+2},"next_node":"old_balanced","death_risk":0.0}]),
        ("midlife_coach", 45,
         "Con 45 entrenas a un joven talento que podría ser campeón.",
         [{"text":"Le dedicas todo tu tiempo",
           "stat_delta":{"carisma":+3,"salud":-1},"next_node":"old_mentor","death_risk":0.0},
          {"text":"Equilibras tu vida y la de él",
           "stat_delta":{"salud":+2,"carisma":+1},"next_node":"old_family","death_risk":0.0}]),
        ("midlife_activist", 45,
         "Con 45 lideras un movimiento social. El régimen empieza a amenazarte.",
         [{"text":"Sigues adelante, públicamente",
           "stat_delta":{"valentia":+4,"carisma":+2},"next_node":"death_martyr","death_risk":0.4},
          {"text":"Operas desde las sombras",
           "stat_delta":{"inteligencia":+3},"next_node":"old_statesman","death_risk":0.05}]),
        ("midlife_thinker", 45,
         "Tu obra es reconocida. Un movimiento quiere usarla como propaganda.",
         [{"text":"Rechazas cualquier uso ideológico",
           "stat_delta":{"inteligencia":+2,"carisma":+2},"next_node":"old_sage","death_risk":0.0},
          {"text":"La cedes a cambio de financiación",
           "stat_delta":{"dinero":+3,"inteligencia":-1},"next_node":"old_tycoon","death_risk":0.0}]),
        ("midlife_hermit", 45,
         "Con 45, un periodista encuentra tus escritos y los publica sin permiso.",
         [{"text":"Los demandas y recuperas el control",
           "stat_delta":{"inteligencia":+2},"next_node":"old_sage","death_risk":0.0},
          {"text":"Los abrazas; el mundo merece leerlos",
           "stat_delta":{"carisma":+3,"suerte":+2},"next_node":"old_icon","death_risk":0.0}]),
        ("midlife_corporate", 45,
         "Con 45 descubres que tu empresa contamina ilegalmente.",
         [{"text":"Lo denuncias públicamente",
           "stat_delta":{"carisma":+3,"inteligencia":+1},"next_node":"old_redeemed","death_risk":0.05},
          {"text":"Cierras los ojos; hay hipoteca que pagar",
           "stat_delta":{"dinero":+2},"next_node":"old_quiet","death_risk":0.0}]),
    ]
    for args in midlife:
        add(*args)

    finals = {
        "old_icon":         (75, "El Ícono",       "Mueres siendo una leyenda viva. El mundo nunca te olvida."),
        "old_sage":         (80, "El Sabio",        "Mueres en paz, rodeado de libros y admiradores."),
        "old_judge":        (78, "El Justo",        "Mueres habiendo impartido justicia durante décadas."),
        "old_family":       (82, "El Patriarca",    "Mueres rodeado de hijos y nietos. Una vida plena."),
        "old_redeemed":     (72, "El Redimido",     "Mueres habiendo expiado tus errores. La paz llega al final."),
        "old_ghost":        (70, "El Fantasma",     "Mueres con una identidad falsa en un país lejano."),
        "old_retired_rich": (85, "El Rentista",     "Mueres rico y tranquilo, sin arrepentimientos."),
        "old_tycoon":       (68, "El Magnate",      "Mueres en la cima del mundo, solo y lleno de logros."),
        "old_statesman":    (77, "El Estadista",    "Mueres recordado como arquitecto de una nación."),
        "old_quiet":        (79, "El Tranquilo",    "Mueres en silencio, sin ruido ni gloria."),
        "old_balanced":     (84, "El Equilibrado",  "Mueres en armonía. Viviste bien, sin excesos."),
        "old_philanthropist":(76,"El Filántropo",   "Mueres sabiendo que tu dinero cambió miles de vidas."),
        "old_mentor":       (74, "El Mentor",       "Mueres viendo a tus pupilos triunfar."),
        "old_thinker":      (81, "El Pensador",     "Mueres dejando ideas que se estudiarán por siglos."),
    }
    for nid, (age, title, msg) in finals.items():
        add(nid, age, f"Llegas a los {age} años. Es el momento de partir.",
            [], True, msg, title)

    deaths = {
        "death_overdose": (27, "El Maldito",      "Mueres a los 27 de sobredosis. El Club de los 27 te reclama."),
        "death_fugitive": (34, "El Fugitivo",     "Mueres a los 34 en un tiroteo policial en tierra extranjera."),
        "death_scandal":  (47, "El Caído",        "Mueres a los 47 en circunstancias misteriosas."),
        "death_hero":     (46, "El Héroe Oscuro", "Mueres a los 46 en misión clasificada. Tu nombre nunca se sabrá."),
        "death_illness":  (46, "El Explorador",   "Mueres a los 46 de una enfermedad sin nombre en la selva."),
        "death_heart":    (48, "El Trabajador",   "Mueres a los 48 de infarto. El trabajo fue tu vida y tu muerte."),
        "death_martyr":   (46, "El Mártir",       "Mueres a los 46 por defender tus ideales. Tu causa sobrevive."),
    }
    for nid, (age, title, msg) in deaths.items():
        add(nid, age, f"El destino llega a los {age} años...", [], True, msg, title)

    return nodes


# ══════════════════════════════════════════════════════════════════════════════
#  PERSONAJE
# ══════════════════════════════════════════════════════════════════════════════

class Character:
    DEFAULT_STATS = {"salud":10,"inteligencia":5,"carisma":5,
                     "valentia":5,"dinero":5,"suerte":5}
    def __init__(self, name):
        self.name    = name
        self.age     = 0
        self.stats   = dict(self.DEFAULT_STATS)
        self.history = []

    def apply_delta(self, delta):
        for k, v in delta.items():
            if k in self.stats:
                self.stats[k] = max(0, min(20, self.stats[k] + v))

    def record(self, node_id, label, choice, color="#2196F3"):
        self.history.append({"node_id":node_id,"label":label,
                             "choice":choice,"age":self.age,"color":color})


# ══════════════════════════════════════════════════════════════════════════════
#  PERSISTENCIA
# ══════════════════════════════════════════════════════════════════════════════

class AchievementDB:
    def __init__(self, path=DB_PATH):
        self.path = path
        with sqlite3.connect(path) as c:
            c.execute("CREATE TABLE IF NOT EXISTS achievements "
                      "(title TEXT PRIMARY KEY, unlocked_at TEXT NOT NULL)")

    def unlock(self, title):
        with sqlite3.connect(self.path) as c:
            cur = c.execute("INSERT OR IGNORE INTO achievements VALUES (?,?)",
                            (title, datetime.now().isoformat(timespec="seconds")))
            return cur.rowcount > 0

    def get_unlocked(self):
        with sqlite3.connect(self.path) as c:
            rows = c.execute("SELECT title,unlocked_at FROM achievements "
                             "ORDER BY unlocked_at").fetchall()
        return {r[0]: r[1] for r in rows}


# ══════════════════════════════════════════════════════════════════════════════
#  GRAFO
# ══════════════════════════════════════════════════════════════════════════════

def render_graph(character, out=GRAPH_OUT):
    if not HAS_MPL: return None
    hist = character.history
    if not hist: return None

    fig, ax = plt.subplots(figsize=(max(10, len(hist)*2), 5.5))
    ax.set_facecolor("#0D0D1A"); fig.patch.set_facecolor("#0D0D1A")

    ages    = [h["age"] for h in hist]
    span    = max(max(ages)-min(ages), 1)
    pos     = {}
    for i, h in enumerate(hist):
        pos[h["node_id"]] = (i, (h["age"]-min(ages))/span)

    if HAS_NX:
        G = nx.DiGraph()
        for h in hist: G.add_node(h["node_id"])
        for i in range(len(hist)-1):
            G.add_edge(hist[i]["node_id"], hist[i+1]["node_id"],
                       label=hist[i]["choice"][:28])
        colors = [h["color"] for h in hist]
        nx.draw_networkx_nodes(G, pos, ax=ax, node_color=colors,
                               node_size=900, alpha=0.93)
        lbl = {h["node_id"]: f"{h['label'][:18]}\n({h['age']}a)" for h in hist}
        nx.draw_networkx_labels(G, pos, labels=lbl, ax=ax,
                                font_size=6.5, font_color="white")
        nx.draw_networkx_edges(G, pos, ax=ax, edge_color="#6060A0",
                               arrows=True, arrowsize=20,
                               connectionstyle="arc3,rad=0.12", width=1.8)
        el = nx.get_edge_attributes(G, "label")
        nx.draw_networkx_edge_labels(G, pos, edge_labels=el, ax=ax,
                                     font_size=5.5, font_color="#9090C0",
                                     bbox=dict(boxstyle="round,pad=0.2",
                                               fc="#1A1A2E", ec="none", alpha=0.8))
    else:
        for i in range(len(hist)-1):
            n1, n2 = hist[i]["node_id"], hist[i+1]["node_id"]
            x1,y1 = pos[n1]; x2,y2 = pos[n2]
            ax.annotate("", xy=(x2,y2), xytext=(x1,y1),
                        arrowprops=dict(arrowstyle="->",color="#6060A0",lw=1.8))
        for h in hist:
            x,y = pos[h["node_id"]]
            c   = plt.Circle((x,y),0.05,color=h["color"],zorder=3)
            ax.add_patch(c)
            ax.text(x,y-0.09,f"{h['label'][:16]}\n({h['age']}a)",
                    color="white",fontsize=6,ha="center",va="top")
        ax.set_xlim(-0.5,len(hist)-0.5); ax.set_ylim(-0.2,1.2)

    legend = [mpatches.Patch(color="#4CAF50",label="Inicio"),
              mpatches.Patch(color="#2196F3",label="Decisión"),
              mpatches.Patch(color="#F44336",label="Muerte prematura"),
              mpatches.Patch(color="#9C27B0",label="Vejez")]
    ax.legend(handles=legend,loc="lower right",facecolor="#1A1A2E",
              labelcolor="white",fontsize=8)
    ax.set_title(f"Camino de vida — {character.name}",
                 color="white",fontsize=13,pad=10)
    ax.axis("off")
    plt.tight_layout()
    plt.savefig(out, dpi=110, bbox_inches="tight", facecolor="#0D0D1A")
    plt.close(fig)
    return out


# ══════════════════════════════════════════════════════════════════════════════
#  HELPERS DE DIBUJO  (pygame)
# ══════════════════════════════════════════════════════════════════════════════

def lerp_color(a, b, t):
    return tuple(int(a[i]+(b[i]-a[i])*t) for i in range(3))

def draw_rect_alpha(surf, color, rect, alpha=180, radius=8):
    s = pygame.Surface((rect[2], rect[3]), pygame.SRCALPHA)
    pygame.draw.rect(s, (*color, alpha), (0,0,rect[2],rect[3]), border_radius=radius)
    surf.blit(s, (rect[0], rect[1]))

def draw_border_glow(surf, rect, color, radius=8, width=2):
    pygame.draw.rect(surf, color, rect, width, border_radius=radius)

def wrap_text(font, text, max_w):
    words  = text.split()
    lines  = []
    cur    = []
    for w in words:
        test = " ".join(cur + [w])
        if font.size(test)[0] <= max_w:
            cur.append(w)
        else:
            if cur: lines.append(" ".join(cur))
            cur = [w]
    if cur: lines.append(" ".join(cur))
    return lines

def draw_text_wrapped(surf, font, text, x, y, max_w, color, line_h=None, max_y=None):
    lh = line_h or font.get_linesize() + 2
    for line in wrap_text(font, text, max_w):
        if max_y is not None and y + lh > max_y:
            surf.blit(font.render('  ▼', True, (90,90,130)), (x, y - lh)); break
        surf.blit(font.render(line, True, color), (x, y)); y += lh
    return y

# ══════════════════════════════════════════════════════════════════════════════
#  PARTÍCULAS DE FONDO
# ══════════════════════════════════════════════════════════════════════════════

class Particle:
    def __init__(self):
        self.reset()

    def reset(self):
        self.x   = random.uniform(0, W)
        self.y   = random.uniform(0, H)
        self.r   = random.uniform(0.4, 1.6)
        self.spd = random.uniform(0.1, 0.35)
        self.a   = random.randint(30, 120)
        self.col = random.choice([(140,90,255),(80,200,255),(212,175,55)])

    def update(self):
        self.y -= self.spd
        if self.y < -5: self.reset(); self.y = H + 5

    def draw(self, surf):
        s = pygame.Surface((4,4), pygame.SRCALPHA)
        pygame.draw.circle(s, (*self.col, self.a), (2,2), int(self.r)+1)
        surf.blit(s, (int(self.x)-2, int(self.y)-2))

PARTICLES = [Particle() for _ in range(120)]

# ══════════════════════════════════════════════════════════════════════════════
#  BOTÓN
# ══════════════════════════════════════════════════════════════════════════════

class Button:
    def __init__(self, rect, text, font, color=None, text_color=None,
                 border_color=None, radius=10):
        self.rect         = pygame.Rect(rect)
        self.text         = text
        self.font         = font
        self.color        = color or C["panel2"]
        self.text_color   = text_color or C["text"]
        self.border_color = border_color or C["border"]
        self.radius       = radius
        self.hovered      = False
        self._hover_t     = 0.0

    def update(self, mx, my):
        self.hovered = self.rect.collidepoint(mx, my)
        target = 1.0 if self.hovered else 0.0
        self._hover_t += (target - self._hover_t) * 0.18

    def draw(self, surf):
        t = self._hover_t
        bg  = lerp_color(self.color, C["panel"], t)
        bdr = lerp_color(self.border_color, C["border_hi"], t)
        draw_rect_alpha(surf, bg, self.rect, alpha=int(200+30*t), radius=self.radius)
        draw_border_glow(surf, self.rect, bdr, self.radius, 2)
        lbl = self.font.render(self.text, True,
                               lerp_color(self.text_color, C["white"], t))
        surf.blit(lbl, lbl.get_rect(center=self.rect.center))

    def clicked(self, event):
        return (event.type == pygame.MOUSEBUTTONDOWN and
                event.button == 1 and self.rect.collidepoint(event.pos))


# ══════════════════════════════════════════════════════════════════════════════
#  SCREENS  (cada pantalla es una clase con .run())
# ══════════════════════════════════════════════════════════════════════════════

class Screen:
    """Clase base."""
    def __init__(self, app):
        self.app  = app
        self.surf = app.screen

    def draw_bg(self):
        self.surf.fill(C["bg"])
        for p in PARTICLES:
            p.update(); p.draw(self.surf)

    def draw_title_bar(self, text):
        # Línea decorativa superior
        pygame.draw.rect(self.surf, C["accent"], (0, 0, W, 3))
        t = self.app.f_title.render(text, True, C["gold"])
        self.surf.blit(t, t.get_rect(centerx=W//2, top=22))
        # Separador
        pygame.draw.line(self.surf, C["border"], (60, 70), (W-60, 70), 1)

    def draw_footer(self, text=""):
        pygame.draw.line(self.surf, C["border"], (60, H-52), (W-60, H-52), 1)
        if text:
            lbl = self.app.f_small.render(text, True, C["text_dim"])
            self.surf.blit(lbl, lbl.get_rect(centerx=W//2, bottom=H-14))

    def handle_events(self):
        for e in pygame.event.get():
            if e.type == pygame.QUIT: pygame.quit(); sys.exit()
            yield e


# ─────────────────────────────────────────────────────────────────────────────
class ScreenMenu(Screen):
    def __init__(self, app):
        super().__init__(app)
        bw, bh = 340, 58
        cx = W // 2
        self.buttons = [
            Button((cx-bw//2, 260, bw, bh), "⚔  JUGAR",
                   app.f_btn, C["panel2"], C["gold"], C["gold_dim"], 12),
            Button((cx-bw//2, 335, bw, bh), "🏆  LOGROS",
                   app.f_btn, C["panel2"], C["accent2"], C["border"], 12),
            Button((cx-bw//2, 410, bw, bh), "🌍  DIFICULTAD  —  " + app.difficulty,
                   app.f_btn, C["panel2"], C["text"], C["border"], 12),
            Button((cx-bw//2, 485, bw, bh), "✖  SALIR",
                   app.f_btn, C["panel2"], C["red"], C["red_dim"], 12),
        ]
        self._pulse = 0.0

    def run(self):
        clock = pygame.time.Clock()
        while True:
            mx, my = pygame.mouse.get_pos()
            self._pulse = (self._pulse + 0.03) % (2 * math.pi)

            for btn in self.buttons: btn.update(mx, my)

            for e in self.handle_events():
                if self.buttons[0].clicked(e): return "play"
                if self.buttons[1].clicked(e): return "achievements"
                if self.buttons[2].clicked(e): return "difficulty"
                if self.buttons[3].clicked(e): pygame.quit(); sys.exit()

            # ── Dibujo ────────────────────────────────────────────────────
            self.draw_bg()

            # Título grande con efecto glow
            glow_a = int(120 + 60 * math.sin(self._pulse))
            for off in [(2,2),(-2,-2),(2,-2),(-2,2)]:
                gs = self.app.f_huge.render("LIFE SIM", True, C["accent"])
                gs.set_alpha(glow_a // 3)
                self.surf.blit(gs, gs.get_rect(centerx=W//2+off[0], centery=155+off[1]))
            t = self.app.f_huge.render("LIFE SIM", True, C["gold"])
            self.surf.blit(t, t.get_rect(centerx=W//2, centery=155))

            sub = self.app.f_med.render("Simulador de Vida", True, C["text_dim"])
            self.surf.blit(sub, sub.get_rect(centerx=W//2, centery=200))

            # Botones
            for btn in self.buttons: btn.draw(self.surf)

            # Versión
            ver = self.app.f_small.render("v2.0  GUI Edition", True, C["text_dim"])
            self.surf.blit(ver, (W-160, H-30))

            pygame.display.flip()
            clock.tick(FPS)


# ─────────────────────────────────────────────────────────────────────────────
class ScreenDifficulty(Screen):
    def __init__(self, app):
        super().__init__(app)
        regions = list(DIFFICULTY_SETTINGS.keys())
        self.buttons = []
        bw, bh = 480, 56
        for i, r in enumerate(regions):
            info = DIFFICULTY_SETTINGS[r]
            lbl  = f"{info['emoji']}  {r:<16}  —  {info['label']}"
            sel  = (r == app.difficulty)
            bc   = C["gold_dim"] if sel else C["border"]
            tc   = C["gold"]     if sel else C["text"]
            btn  = Button((W//2-bw//2, 130 + i*72, bw, bh), lbl,
                          app.f_btn, C["panel2"], tc, bc, 12)
            btn._region = r
            self.buttons.append(btn)
        self.back = Button((40, H-70, 160, 44), "← VOLVER",
                           app.f_small, C["panel2"], C["text_dim"], C["border"], 8)

    def run(self):
        clock = pygame.time.Clock()
        while True:
            mx, my = pygame.mouse.get_pos()
            for btn in self.buttons: btn.update(mx, my)
            self.back.update(mx, my)

            for e in self.handle_events():
                for btn in self.buttons:
                    if btn.clicked(e):
                        self.app.difficulty = btn._region
                        return "menu"
                if self.back.clicked(e): return "menu"

            self.draw_bg()
            self.draw_title_bar("SELECCIONAR REGIÓN")

            # Sub-header
            sh = self.app.f_med.render("Cada región determina el riesgo de muerte", True, C["text_dim"])
            self.surf.blit(sh, sh.get_rect(centerx=W//2, top=82))

            for btn in self.buttons:
                # Marcar seleccionada
                if btn._region == self.app.difficulty:
                    draw_rect_alpha(self.surf, C["accent"], btn.rect, 30, 12)
                btn.draw(self.surf)
                # Mostrar multiplicador
                info = DIFFICULTY_SETTINGS[btn._region]
                rm   = self.app.f_small.render(f"×{info['risk_mult']:.1f}", True, C["accent2"])
                self.surf.blit(rm, rm.get_rect(right=W//2+260, centery=btn.rect.centery))

            self.back.draw(self.surf)
            self.draw_footer("Dificultad afecta la probabilidad de muerte súbita")
            pygame.display.flip()
            clock.tick(FPS)


# ─────────────────────────────────────────────────────────────────────────────
class ScreenAchievements(Screen):
    def __init__(self, app):
        super().__init__(app)
        self.back   = Button((40, H-70, 160, 44), "← VOLVER",
                             app.f_small, C["panel2"], C["text_dim"], C["border"], 8)
        self.scroll = 0

    def run(self):
        clock   = pygame.time.Clock()
        unlocked = self.app.db.get_unlocked()
        while True:
            mx, my = pygame.mouse.get_pos()
            self.back.update(mx, my)

            for e in self.handle_events():
                if self.back.clicked(e): return "menu"
                if e.type == pygame.MOUSEWHEEL:
                    self.scroll = max(0, self.scroll - e.y * 24)

            self.draw_bg()
            self.draw_title_bar("LOGROS Y TÍTULOS")

            got   = len(unlocked)
            total = len(ALL_TITLES)
            prog  = self.app.f_med.render(f"{got} / {total} títulos desbloqueados",
                                          True, C["accent2"])
            self.surf.blit(prog, prog.get_rect(centerx=W//2, top=82))

            # Barra de progreso
            bx, by, bw, bh2 = W//2-200, 112, 400, 10
            draw_rect_alpha(self.surf, C["panel2"], (bx,by,bw,bh2), 200, 5)
            fill_w = int(bw * got / total)
            if fill_w > 0:
                draw_rect_alpha(self.surf, C["accent"], (bx,by,fill_w,bh2), 255, 5)

            # Lista (con scroll)
            row_h  = 46
            list_y = 140
            vis_h  = H - 160
            clip   = pygame.Rect(0, list_y, W, vis_h)
            self.surf.set_clip(clip)

            for i, title in enumerate(ALL_TITLES):
                y = list_y + i * row_h - self.scroll
                if y + row_h < list_y or y > list_y + vis_h: continue

                if title in unlocked:
                    draw_rect_alpha(self.surf, C["panel2"], (W//2-320, y, 640, row_h-4), 200, 8)
                    draw_border_glow(self.surf, pygame.Rect(W//2-320, y, 640, row_h-4),
                                     C["gold_dim"], 8, 1)
                    icon = self.app.f_med.render("✦", True, C["gold"])
                    self.surf.blit(icon, (W//2-305, y+10))
                    lbl = self.app.f_med.render(title, True, C["gold"])
                    self.surf.blit(lbl, (W//2-275, y+10))
                    dt  = self.app.f_small.render(unlocked[title][:16], True, C["text_dim"])
                    self.surf.blit(dt, dt.get_rect(right=W//2+315, centery=y+row_h//2))
                else:
                    draw_rect_alpha(self.surf, C["panel"], (W//2-320, y, 640, row_h-4), 130, 8)
                    lbl = self.app.f_med.render("??? — Logro bloqueado", True, C["border"])
                    self.surf.blit(lbl, lbl.get_rect(centerx=W//2, centery=y+row_h//2))

            self.surf.set_clip(None)
            self.back.draw(self.surf)
            self.draw_footer("Scroll para ver más")
            pygame.display.flip()
            clock.tick(FPS)


# ─────────────────────────────────────────────────────────────────────────────
class ScreenNameInput(Screen):
    def __init__(self, app):
        super().__init__(app)
        self.name   = ""
        self.cursor = True
        self._ct    = 0
        self.start  = Button((W//2-120, 400, 240, 54), "COMENZAR →",
                             app.f_btn, C["panel2"], C["gold"], C["gold_dim"], 12)
        self.back   = Button((40, H-70, 160, 44), "← VOLVER",
                             app.f_small, C["panel2"], C["text_dim"], C["border"], 8)

    def run(self):
        clock = pygame.time.Clock()
        while True:
            mx, my = pygame.mouse.get_pos()
            self._ct += 1
            if self._ct % 30 == 0: self.cursor = not self.cursor
            self.start.update(mx, my); self.back.update(mx, my)

            for e in self.handle_events():
                if e.type == pygame.KEYDOWN:
                    if e.key == pygame.K_BACKSPACE:
                        self.name = self.name[:-1]
                    elif e.key == pygame.K_RETURN and self.name.strip():
                        return self.name.strip()
                    elif len(self.name) < 24:
                        self.name += e.unicode
                if self.start.clicked(e) and self.name.strip():
                    return self.name.strip()
                if self.back.clicked(e): return None

            self.draw_bg()
            self.draw_title_bar("NUEVA VIDA")

            prompt = self.app.f_med.render("¿Cómo se llamará tu personaje?", True, C["text_dim"])
            self.surf.blit(prompt, prompt.get_rect(centerx=W//2, top=200))

            # Caja de texto
            bx, by, bw, bh = W//2-220, 270, 440, 60
            draw_rect_alpha(self.surf, C["panel2"], (bx,by,bw,bh), 220, 10)
            draw_border_glow(self.surf, pygame.Rect(bx,by,bw,bh), C["border_hi"], 10, 2)
            display = self.name + ("|" if self.cursor else " ")
            t = self.app.f_title.render(display, True, C["text"])
            self.surf.blit(t, t.get_rect(center=(W//2, by+bh//2)))

            # Región
            info = DIFFICULTY_SETTINGS[self.app.difficulty]
            reg  = self.app.f_small.render(
                f"Región: {self.app.difficulty} ({info['label']})", True, C["accent2"])
            self.surf.blit(reg, reg.get_rect(centerx=W//2, top=350))

            self.start.draw(self.surf)
            self.back.draw(self.surf)
            pygame.display.flip()
            clock.tick(FPS)


# ─────────────────────────────────────────────────────────────────────────────
class ScreenGame(Screen):
    """Pantalla principal de juego — muestra situación, stats y botones de elección."""

    STAT_ICONS = {
        "salud":"❤️", "inteligencia":"🧠", "carisma":"🗣️",
        "valentia":"⚔️", "dinero":"💰", "suerte":"🎲"
    }

    def __init__(self, app, character, tree, risk_mult):
        super().__init__(app)
        self.char      = character
        self.tree      = tree
        self.risk_mult = risk_mult
        self.current   = "start"
        self.result    = None   # "death_premature" | "death_old" | "victory"
        self.title_earned = ""
        self.death_msg    = ""
        self._anim_alpha  = 0   # fade-in pantalla
        self._choice_btns = []
        self._build_choice_buttons(tree["start"])
        self._msg_log     = []  # mensajes de narrative recientes
        self._shake       = 0   # frames de shake en muerte

    def _build_choice_buttons(self, node):
        self._choice_btns = []
        bw, bh = 560, 70
        # Posición vertical: justo encima del panel de estadísticas, que ahora está en H-140
        start_y = H - 160 - (len(node.options) * (bh + 10))
        for i, opt in enumerate(node.options):
            y   = start_y + i * (bh + 10)
            btn = Button((W//2-bw//2, y, bw, bh), opt["text"],
                         self.app.f_btn, C["panel2"], C["text"], C["border"], 12)
            btn._opt = opt
            self._choice_btns.append(btn)

    def run(self):
        clock = pygame.time.Clock()
        fade  = 255

        while True:
            mx, my = pygame.mouse.get_pos()
            node   = self.tree.get(self.current)

            if node and node.is_terminal:
                return self._handle_terminal(node)

            # Nodo sin opciones: extraer redirect del texto o terminar
            if node and not node.is_terminal and not node.options:
                import re as _re_safe
                _rm = _re_safe.search(r'->\s*[Vv]e\s+a\s+la\s+(\d+)', node.situation)
                if _rm:
                    _n = int(_rm.group(1))
                    for _cand in [f"q_{_n}", f"death_{_n}", f"old_{_n}"]:
                        if _cand in self.tree:
                            self.current = _cand
                            node = self.tree.get(self.current)
                            if node and not node.is_terminal:
                                self._build_choice_buttons(node)
                            break
                    else:
                        self._random_death()
                else:
                    self._random_death()

            for btn in self._choice_btns: btn.update(mx, my)

            for e in self.handle_events():
                for i, btn in enumerate(self._choice_btns):
                    if btn.clicked(e) and node:
                        self._process_choice(node, btn._opt)
                        node = self.tree.get(self.current)
                        if node and not node.is_terminal:
                            self._build_choice_buttons(node)

            if self._shake > 0:
                self._shake -= 1

            # ── Dibujo ────────────────────────────────────────────────────
            ox = random.randint(-3,3) if self._shake > 0 else 0
            self.surf.fill(C["bg"])
            for p in PARTICLES: p.update(); p.draw(self.surf)

            if node:
                self._draw_game(node, ox)

            # Fade-in
            if fade > 0:
                s = pygame.Surface((W,H)); s.set_alpha(fade); s.fill((0,0,0))
                self.surf.blit(s,(0,0)); fade = max(0, fade-12)

            pygame.display.flip()
            clock.tick(FPS)

    def _process_choice(self, node, opt):
        color = "#4CAF50" if self.current == "start" else "#2196F3"
        self.char.record(self.current,
                         node.situation[:22]+"…" if len(node.situation)>22 else node.situation,
                         opt["text"], color)
        self.char.apply_delta(opt["stat_delta"])
        self.char.age = node.age

        # Risk check
        eff = opt["death_risk"] * self.risk_mult
        if eff > 0 and random.random() < eff:
            self._random_death()
            return

        self.current = opt["next_node"]
        next_node = self.tree.get(self.current)
        if next_node: self.char.age = next_node.age
        self._shake = 8
        self._msg_log = [f"→ {opt['text'][:55]}"]

    def _random_death(self):
        self.char.record("death_random","Muerte inesperada","— Fin —","#F44336")
        self.title_earned = "El Temerario"
        self.death_msg    = (f"Mueres a los {self.char.age} años por las consecuencias "
                             f"de tu decisión. La dificultad '{self.app.difficulty}' cobró su precio.")
        self.result       = "death_premature"
        self.current      = "__dead__"

    def _handle_terminal(self, node):
        is_pre = node.node_id.startswith("death_")
        color  = "#F44336" if is_pre else "#9C27B0"
        self.char.age = node.age
        self.char.record(node.node_id, node.death_msg[:22], "— Fin —", color)
        self.title_earned = node.title
        self.death_msg    = node.death_msg
        self.result       = "death_premature" if is_pre else "death_old"
        return ScreenDeath(self.app, self.char, self.title_earned,
                           self.death_msg, self.result).run()

    def _draw_game(self, node, ox=0):
        # ── Panel principal ────────────────────────────────────────────
        panel_rect = (ox+40, 10, W-80, H-240)
        draw_rect_alpha(self.surf, C["panel"], panel_rect, 200, 14)
        draw_border_glow(self.surf, pygame.Rect(*panel_rect), C["border"], 14, 1)

        # Header: nombre + edad + región
        header_r = (ox+40, 10, W-80, 50)
        draw_rect_alpha(self.surf, C["panel2"], header_r, 230, 14)
        name_lbl = self.app.f_med.render(
            f"{self.char.name}  ·  {self.char.age} años", True, C["gold"])
        self.surf.blit(name_lbl, (ox+60, 22))
        diff_lbl = self.app.f_small.render(
            f"{self.app.difficulty}  ({DIFFICULTY_SETTINGS[self.app.difficulty]['label']})",
            True, C["text_dim"])
        self.surf.blit(diff_lbl, diff_lbl.get_rect(right=ox+W-60, centery=35))

        # Situación
        _btn_top = H - 160 - (len(node.options) * 80) - 10
        sit_y = draw_text_wrapped(self.surf, self.app.f_sit,
                                  node.situation, ox+60, 80, W-160,
                                  C["text"], line_h=30, max_y=_btn_top - 10)

        # Últimas acciones
        if self._msg_log:
            log_lbl = self.app.f_small.render(self._msg_log[-1], True, C["text_dim"])
            self.surf.blit(log_lbl, (ox+62, max(sit_y+10, 140)))

        # ── Stats (MEJORADOS) ──────────────────────────────────────────
        self._draw_stats(ox)

        # ── Botones de elección ────────────────────────────────────────
        for i, btn in enumerate(self._choice_btns):
            r = btn.rect.copy()
            r.x += ox
            # Delta hint
            opt    = btn._opt
            deltas = [f"{'+' if v>=0 else ''}{v} {k}" for k,v in opt["stat_delta"].items() if v!=0]
            draw_rect_alpha(self.surf, C["panel2"], r, 210, 12)
            draw_border_glow(self.surf, r,
                             C["border_hi"] if btn.hovered else C["border"], 12, 2)
            _tc=lerp_color(C["text"],C["white"],btn._hover_t)
            _ws=btn.text.split(); _l1,_l2,_sw=[],[],False
            for _w in _ws:
                if not _sw and self.app.f_btn.size(" ".join(_l1+[_w]))[0]<r.width-20: _l1.append(_w)
                else: _sw=True; _l2.append(_w)
            if _l2:
                _s1=self.app.f_btn.render(" ".join(_l1),True,_tc)
                _s2=self.app.f_btn.render(" ".join(_l2),True,_tc)
                self.surf.blit(_s1,_s1.get_rect(centerx=r.centerx,centery=r.top+16))
                self.surf.blit(_s2,_s2.get_rect(centerx=r.centerx,centery=r.top+36))
            else:
                _s=self.app.f_btn.render(" ".join(_l1),True,_tc)
                self.surf.blit(_s,_s.get_rect(centerx=r.centerx,centery=r.top+22))
            if deltas:
                dstr = "  ".join(deltas[:4])
                dl   = self.app.f_small.render(dstr, True, C["accent2"])
                self.surf.blit(dl, dl.get_rect(centerx=r.centerx, bottom=r.bottom-8))

            # Riesgo
            eff = opt["death_risk"] * self.risk_mult
            if eff > 0.05:
                col = C["red"] if eff > 0.3 else C["orange"]
                rl  = self.app.f_small.render(f"⚠ {int(eff*100)}% riesgo", True, col)
                self.surf.blit(rl, rl.get_rect(right=r.right-14, centery=r.centery))

    def _draw_stats(self, ox=0):
        stats = self.char.stats
        # Panel de estadísticas: más grande, ubicado justo encima de los botones
        panel_h = 110
        panel_y = H - panel_h - 20
        panel_w = W - 80
        sx = ox + 40
        sy = panel_y
        panel_rect = (sx, sy, panel_w, panel_h)
        draw_rect_alpha(self.surf, C["panel2"], panel_rect, 220, 12)
        draw_border_glow(self.surf, pygame.Rect(panel_rect), C["border_hi"], 12, 1)

        # Organización en dos filas de 3 atributos
        keys = ["salud", "inteligencia", "carisma", "valentia", "dinero", "suerte"]
        titles = {
            "salud": "Salud", "inteligencia": "Inteligencia",
            "carisma": "Carisma", "valentia": "Valentía",
            "dinero": "Dinero", "suerte": "Suerte"
        }

        margin_top = 12
        row_height = (panel_h - margin_top*2) // 2
        col_width = panel_w // 3

        mx, my = pygame.mouse.get_pos()
        mx -= ox   # ajustar por scroll/shake

        for idx, key in enumerate(keys):
            row = idx // 3
            col = idx % 3
            x = sx + col * col_width + 15
            y = sy + margin_top + row * row_height + 4

            v = stats[key]
            color = C["stat_bars"].get(key, C["accent"])
            bar_max_width = col_width - 80
            bar_width = int(bar_max_width * v / 20)

            # Icono y nombre
            icon = self.STAT_ICONS.get(key, "■")
            name = titles.get(key, key.capitalize())
            lbl = self.app.f_small.render(f"{icon} {name}", True, C["text"])
            self.surf.blit(lbl, (x, y))

            # Fondo de la barra
            bar_bg_rect = (x + 5, y + 20, bar_max_width, 16)
            pygame.draw.rect(self.surf, C["border"], bar_bg_rect, border_radius=8)

            # Relleno
            if bar_width > 0:
                bar_fill_rect = (x + 5, y + 20, bar_width, 16)
                # Degradado sutil
                grad = pygame.Surface((bar_width, 16), pygame.SRCALPHA)
                grad.fill(color)
                for i in range(16):
                    alpha = 255 - i * 8
                    pygame.draw.line(grad, (255,255,255, max(0,alpha//3)), (0,i), (bar_width,i))
                self.surf.blit(grad, bar_fill_rect)
                pygame.draw.rect(self.surf, color, bar_fill_rect, 1, border_radius=8)

            # Valor numérico
            val_text = str(v)
            val_surf = self.app.f_med.render(val_text, True, color)
            val_rect = val_surf.get_rect(midleft=(x + bar_max_width + 18, y + 28))
            # Sombra
            shadow = self.app.f_med.render(val_text, True, (0,0,0))
            shadow.set_alpha(120)
            self.surf.blit(shadow, (val_rect.x+1, val_rect.y+1))
            self.surf.blit(val_surf, val_rect)

            # Efecto hover
            mouse_over = (mx > bar_bg_rect[0] and mx < bar_bg_rect[0] + bar_bg_rect[2] and
                          my > bar_bg_rect[1] and my < bar_bg_rect[1] + bar_bg_rect[3])
            if mouse_over:
                glow_rect = (bar_bg_rect[0]-2, bar_bg_rect[1]-2,
                             bar_bg_rect[2]+4, bar_bg_rect[3]+4)
                pygame.draw.rect(self.surf, C["white"], glow_rect, 2, border_radius=10)
                big_val = self.app.f_title.render(str(v), True, color)
                big_rect = big_val.get_rect(center=(x + bar_max_width//2, y + 28))
                bg_rect = big_rect.inflate(20, 10)
                draw_rect_alpha(self.surf, (0,0,0), bg_rect, 180, 8)
                self.surf.blit(big_val, big_rect)


# ─────────────────────────────────────────────────────────────────────────────
class ScreenDeath(Screen):
    def __init__(self, app, char, title, msg, kind):
        super().__init__(app)
        self.char   = char
        self.title  = title
        self.msg    = msg
        self.kind   = kind
        self._alpha = 0
        self._t     = 0

        is_new = app.db.unlock(title)
        self.is_new = is_new

        # Generar grafo
        self.graph_path = render_graph(char, GRAPH_OUT) if HAS_MPL else None
        self.graph_surf = None
        if self.graph_path and os.path.exists(self.graph_path):
            raw = pygame.image.load(self.graph_path)
            gw  = W - 120
            gh  = int(raw.get_height() * gw / raw.get_width())
            gh  = min(gh, 200)
            self.graph_surf = pygame.transform.smoothscale(raw, (gw, gh))

        self.btn_again = Button((W//2-180, H-90, 160, 48), "⚔ OTRA VIDA",
                                app.f_btn, C["panel2"], C["gold"], C["gold_dim"], 12)
        self.btn_menu  = Button((W//2+20,  H-90, 160, 48), "⌂ MENÚ",
                                app.f_btn, C["panel2"], C["text_dim"], C["border"], 12)

    def run(self):
        clock = pygame.time.Clock()
        while True:
            self._t += 1
            self._alpha = min(255, self._alpha + 6)
            mx, my = pygame.mouse.get_pos()
            self.btn_again.update(mx, my); self.btn_menu.update(mx, my)

            for e in self.handle_events():
                if self.btn_again.clicked(e): return "play"
                if self.btn_menu.clicked(e):  return "menu"

            # BG
            self.surf.fill(C["bg"])
            for p in PARTICLES: p.update(); p.draw(self.surf)

            # Overlay de color según tipo de muerte
            overlay_col = (40,0,0) if self.kind == "death_premature" else (20,0,40)
            ov = pygame.Surface((W,H), pygame.SRCALPHA)
            ov.fill((*overlay_col, 80)); self.surf.blit(ov,(0,0))

            # Vignette
            self._draw_vignette()

            # Contenido con fade
            content = pygame.Surface((W,H), pygame.SRCALPHA)
            self._draw_death_content(content)
            content.set_alpha(self._alpha)
            self.surf.blit(content, (0,0))

            self.btn_again.draw(self.surf)
            self.btn_menu.draw(self.surf)
            pygame.display.flip()
            clock.tick(FPS)

    def _draw_vignette(self):
        for r in range(300, 0, -30):
            a = int(80 * (1 - r/300))
            s = pygame.Surface((W,H), pygame.SRCALPHA)
            pygame.draw.rect(s,(0,0,0,a),(0,0,W,H),r)
            self.surf.blit(s,(0,0))

    def _draw_death_content(self, surf):
        is_pre = self.kind == "death_premature"

        # Título
        hdr_col = C["red"] if is_pre else C["accent"]
        hdr_txt = "✝ MUERTE PREMATURA" if is_pre else "☽ FIN DE UNA VIDA"
        hdr = self.app.f_title.render(hdr_txt, True, hdr_col)
        surf.blit(hdr, hdr.get_rect(centerx=W//2, top=30))
        pygame.draw.line(surf, hdr_col, (W//2-200,72),(W//2+200,72),1)

        # Epitafio
        y = 90
        for line in wrap_text(self.app.f_sit, self.msg, W-160):
            lbl = self.app.f_sit.render(line, True, C["text"])
            surf.blit(lbl, lbl.get_rect(centerx=W//2, top=y))
            y += 34

        # Edad
        age_lbl = self.app.f_med.render(f"Edad: {self.char.age} años", True, C["text_dim"])
        surf.blit(age_lbl, age_lbl.get_rect(centerx=W//2, top=y+8))

        # Título ganado
        y += 52
        badge_r = pygame.Rect(W//2-200, y, 400, 54)
        pygame.draw.rect(surf, C["panel2"], badge_r, border_radius=12)
        pygame.draw.rect(surf, C["gold_dim"], badge_r, 2, border_radius=12)
        t1 = self.app.f_small.render("TÍTULO OBTENIDO", True, C["text_dim"])
        surf.blit(t1, t1.get_rect(centerx=W//2, top=y+6))
        t2 = self.app.f_title.render(f"✦ {self.title} ✦", True, C["gold"])
        surf.blit(t2, t2.get_rect(centerx=W//2, top=y+24))
        if self.is_new:
            new_lbl = self.app.f_small.render("🏆 ¡NUEVO LOGRO!", True, C["green"])
            surf.blit(new_lbl, new_lbl.get_rect(centerx=W//2, top=y+60))

        # Stats finales
        y += 115
        stats = self.char.stats
        sx = W//2 - 300
        for i, (k, v) in enumerate(stats.items()):
            col = C["stat_bars"].get(k, C["accent"])
            kl  = self.app.f_small.render(k.capitalize(), True, C["text_dim"])
            surf.blit(kl, (sx + i*100, y))
            vl  = self.app.f_med.render(str(v), True, col)
            surf.blit(vl, vl.get_rect(centerx=sx+i*100+35, top=y+18))

        # Grafo
        if self.graph_surf:
            gy = y + 60
            gx = W//2 - self.graph_surf.get_width()//2
            # Fondo
            pygame.draw.rect(surf, C["panel"],
                             (gx-6,gy-6,self.graph_surf.get_width()+12,
                              self.graph_surf.get_height()+12), border_radius=8)
            surf.blit(self.graph_surf, (gx, gy))
            glbl = self.app.f_small.render("Grafo de decisiones", True, C["text_dim"])
            surf.blit(glbl, glbl.get_rect(centerx=W//2, top=gy-20))
        elif HAS_MPL:
            gl = self.app.f_small.render("(grafo guardado en life_graph.png)", True, C["text_dim"])
            surf.blit(gl, gl.get_rect(centerx=W//2, top=y+60))


# ══════════════════════════════════════════════════════════════════════════════
#  APLICACIÓN PRINCIPAL
# ══════════════════════════════════════════════════════════════════════════════

class App:
    def __init__(self):
        pygame.init()
        pygame.display.set_caption("Life Sim — GUI Edition")
        self.screen     = pygame.display.set_mode((W, H))
        self.db         = AchievementDB(DB_PATH)
        self.difficulty = "Europa"
        self.tree       = build_decision_tree()
        self._load_fonts()

    def _load_fonts(self):
        # Intenta cargar fuentes del sistema; cae a la por defecto
        def f(size, bold=False):
            for name in ["Segoe UI","DejaVu Sans","FreeSans","Arial",""]:
                try:
                    return pygame.font.SysFont(name, size, bold=bold)
                except Exception:
                    pass
            return pygame.font.Font(None, size)

        self.f_huge  = f(72, bold=True)
        self.f_title = f(32, bold=True)
        self.f_med   = f(22)
        self.f_btn   = f(20, bold=True)
        self.f_sit   = f(19)
        self.f_small = f(15)

    def run(self):
        state = "menu"
        while True:
            if state == "menu":
                state = ScreenMenu(self).run()

            elif state == "difficulty":
                state = ScreenDifficulty(self).run()

            elif state == "achievements":
                state = ScreenAchievements(self).run()

            elif state == "play":
                name = ScreenNameInput(self).run()
                if name is None:
                    state = "menu"
                    continue
                # Rebuild tree for the selected difficulty
                self.tree = build_difficulty_tree(self.difficulty)
                char       = Character(name)
                risk_mult  = DIFFICULTY_SETTINGS[self.difficulty]["risk_mult"]
                result     = ScreenGame(self, char, self.tree, risk_mult).run()
                state      = result if result in ("menu","play") else "menu"

            else:
                state = "menu"


# ══════════════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    App().run()